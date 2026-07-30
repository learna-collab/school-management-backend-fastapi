from pathlib import Path

import docx
import fitz


class DocumentReader:
    def read(self, path: str) -> str:
        suffix = Path(path).suffix.lower()

        if suffix == ".docx":
            return self.read_docx(path)

        if suffix == ".pdf":
            return self.read_pdf(path)

        raise ValueError("Unsupported document type")

    # -------------------------------------------------
    # DOCX READER
    # Reads BOTH paragraphs and tables
    # -------------------------------------------------

    def read_docx(self, path: str) -> str:
        document = docx.Document(path)

        parts = []

        # Normal paragraphs
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # Tables
        for table in document.tables:
            for row in table.rows:
                cells = []

                for cell in row.cells:
                    cell_text = " ".join(
                        p.text.strip() for p in cell.paragraphs if p.text.strip()
                    )

                    if cell_text:
                        cells.append(cell_text)

                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    # -------------------------------------------------
    # PDF READER
    # -------------------------------------------------

    def read_pdf(self, path: str) -> str:
        document = fitz.open(path)

        text = ""

        for page in document:
            text += page.get_text()

        return text
